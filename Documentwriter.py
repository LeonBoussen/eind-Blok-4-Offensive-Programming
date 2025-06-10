import os
from openpyxl import Workbook, load_workbook
from docx import Document

# Save to Excel
def save_to_excel(name, phone, email, file_path="fake_contacts.xlsx"):
    if os.path.exists(file_path):
        workbook = load_workbook(file_path)
        sheet = workbook.active
    else:
        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["Name", "Phone", "Email"])

    sheet.append([name, phone, email])
    workbook.save(file_path)
    print(f"Contact saved to Excel: {file_path}")

# Save to Word
def save_to_word(name, phone, email, file_path="fake_contacts.docx"):
    if os.path.exists(file_path):
        doc = Document(file_path)
    else:
        doc = Document()
        doc.add_heading("Contact List", 0)

    doc.add_paragraph(f"Name: {name}\nPhone: {phone}\nEmail: {email}\n")
    doc.save(file_path)
    print(f"Contact saved to Word: {file_path}")

def main():
    # Ask user which format to use
    format_choice = input("Save contact to (excel/word)? ").strip().lower()

    # Collect contact info
    name = input("Enter name: ").strip()
    phone = input("Enter phone number: ").strip()
    email = input("Enter email address: ").strip()

    # Choose format
    if format_choice == "excel":
        save_to_excel(name, phone, email)
    elif format_choice == "word":
        save_to_word(name, phone, email)
    else:
        print("Invalid choice. Please choose 'excel' or 'word'.")